"""Document Understanding Layer — shared foundation for all document parsers."""

from __future__ import annotations

import hashlib
import re
from pathlib import Path

from app.schemas.document import Document, DocumentMetadata

SUPPORTED_TYPES = {".pdf", ".docx"}
EXTRACTOR_VERSION = "1.0.0"


def validate_filetype(filename: str) -> str:
  ext = Path(filename).suffix.lower()
  if ext not in SUPPORTED_TYPES:
    raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_TYPES}")
  return ext.lstrip(".")


def clean_text(text: str) -> str:
  """Normalize whitespace, line breaks, symbols."""
  text = text.replace("\r\n", "\n").replace("\r", "\n")
  text = re.sub(r"[ \t]+", " ", text)
  text = re.sub(r"\n{3,}", "\n\n", text)
  text = re.sub(r"[^\S\n]+", " ", text)
  return text.strip()


def extract_text_from_bytes(content: bytes, filetype: str) -> tuple[str, int]:
  """Extract raw text — Phase 1 will add PyMuPDF + python-docx."""
  if filetype == "pdf":
    return _extract_pdf(content)
  if filetype == "docx":
    return _extract_docx(content)
  raise ValueError(f"Unsupported filetype: {filetype}")


def _extract_pdf(content: bytes) -> tuple[str, int]:
  try:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=content, filetype="pdf")
    pages = [page.get_text() for page in doc]
    return "\n\n".join(pages), len(pages)
  except ImportError:
    # Fallback until pymupdf installed
    from PyPDF2 import PdfReader
    import io

    reader = PdfReader(io.BytesIO(content))
    pages = [p.extract_text() or "" for p in reader.pages]
    return "\n\n".join(pages), len(pages)


def _extract_docx(content: bytes) -> tuple[str, int]:
  try:
    import io

    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs), 1
  except ImportError as exc:
    raise ValueError("python-docx not installed — pip install python-docx") from exc


def build_document(filename: str, content: bytes) -> Document:
  filetype = validate_filetype(filename)
  raw_text, page_count = extract_text_from_bytes(content, filetype)
  cleaned = clean_text(raw_text)
  content_hash = hashlib.sha256(content).hexdigest()

  confidence = 1.0 if len(cleaned) > 100 else 0.5

  return Document(
    filename=filename,
    filetype=filetype,
    pages=page_count,
    language="en",  # language_detector.py will replace in Phase 1
    raw_text=raw_text,
    cleaned_text=cleaned,
    sections={},
    metadata=DocumentMetadata(
      file_size_bytes=len(content),
      page_count=page_count,
      content_hash=content_hash,
      extractor_version=EXTRACTOR_VERSION,
    ),
    confidence=confidence,
  )
